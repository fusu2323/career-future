import base64
import os
import pytest
import json
from pathlib import Path
from unittest.mock import MagicMock, AsyncMock
from openai import APITimeoutError, APIStatusError


# Set fake API key before Settings is ever instantiated
os.environ.setdefault("DEEPSEEK_API_KEY", "test-key-for-unit-tests")


@pytest.fixture
def jobs_cleaned_path():
    return Path("data/processed/jobs_cleaned.json")


@pytest.fixture
def vector_db_path():
    return Path("data/vector_db")


@pytest.fixture
def collection(vector_db_path):
    import chromadb
    client = chromadb.PersistentClient(path=str(vector_db_path))
    return client.get_collection("job_postings")


@pytest.fixture
def sample_job_record(jobs_cleaned_path):
    with open(jobs_cleaned_path, encoding="utf-8") as f:
        jobs = json.load(f)
    return jobs[0]


# =============================================================================
# Phase 05 LLM Service Fixtures
# =============================================================================


@pytest.fixture
def mock_deepseek_client():
    """Return a MagicMock that simulates a successful DeepSeek API response."""
    client = MagicMock()
    mock_message = MagicMock()
    mock_message.content = '{"skill":"test","level":5}'
    mock_message.role = "assistant"
    mock_choice = MagicMock()
    mock_choice.message = mock_message
    mock_completion = MagicMock()
    mock_completion.choices = [mock_choice]
    mock_completion.model = "deepseek-chat"
    client.chat.completions.create.return_value = mock_completion
    return client


@pytest.fixture
def mock_deepseek_client_json_fail():
    """Return a client that returns non-JSON on first call, valid JSON on second."""
    client = MagicMock()

    # Fail response
    fail_message = MagicMock()
    fail_message.content = "not json"
    fail_message.role = "assistant"
    fail_choice = MagicMock()
    fail_choice.message = fail_message
    mock_fail = MagicMock()
    mock_fail.choices = [fail_choice]
    mock_fail.model = "deepseek-chat"

    # Success response
    success_message = MagicMock()
    success_message.content = '{"skill":"test","level":5}'
    success_message.role = "assistant"
    success_choice = MagicMock()
    success_choice.message = success_message
    mock_success = MagicMock()
    mock_success.choices = [success_choice]
    mock_success.model = "deepseek-chat"

    client.chat.completions.create.side_effect = [mock_fail, mock_success]
    return client


@pytest.fixture
def mock_timeout_client():
    """Return a client whose chat.completions.create raises APITimeoutError."""
    client = MagicMock()
    client.chat.completions.create.side_effect = APITimeoutError("timed out")
    return client


@pytest.fixture
def mock_500_client():
    """Return a client that raises APIStatusError(500) on first call, succeeds on second."""
    client = MagicMock()

    mock_response_500 = MagicMock()
    mock_response_500.status_code = 500

    success_message = MagicMock()
    success_message.content = '{"skill":"test","level":5}'
    success_message.role = "assistant"
    success_choice = MagicMock()
    success_choice.message = success_message
    mock_success = MagicMock()
    mock_success.choices = [success_choice]
    mock_success.model = "deepseek-chat"

    client.chat.completions.create.side_effect = [
        APIStatusError("server error", response=mock_response_500, body=None),
        mock_success,
    ]
    return client


@pytest.fixture
def mock_400_client():
    """Return a client that raises APIStatusError(400) on every call (should NOT retry)."""
    client = MagicMock()
    mock_response_400 = MagicMock()
    mock_response_400.status_code = 400
    client.chat.completions.create.side_effect = APIStatusError(
        "bad request", response=mock_response_400, body=None
    )
    return client


# =============================================================================
# Phase 06 Resume Parsing Fixtures
# =============================================================================


# Minimal valid PDF (ASCII content - LLM client is mocked so content doesn't matter)
_PDF_ASCII = (
    b"%PDF-1.4\n"
    b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n"
    b"2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n"
    b"3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
    b"/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>\nendobj\n"
    b"4 0 obj\n<< /Length 44 >>\nstream\n"
    b"BT\n/F1 12 Tf\n100 700 Td\n"
    b"(Test Resume) Tj\n"
    b"ET\nendstream\nendobj\n"
    b"5 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\n"
    b"xref\n0 6\n"
    b"0000000000 65535 f\n"
    b"0000000009 00000 n\n"
    b"0000000058 00000 n\n"
    b"0000000115 00000 n\n"
    b"0000000266 00000 n\n"
    b"0000000356 00000 n\n"
    b"trailer\n<< /Size 6 /Root 1 0 R >>\n"
    b"startxref\n"
    b"430\n"
    b"%%EOF"
)


@pytest.fixture
def sample_pdf_bytes():
    """Return minimal PDF bytes for testing (single page)."""
    return _PDF_ASCII


@pytest.fixture
def sample_docx_bytes():
    """Return minimal DOCX bytes for testing."""
    # python-docx reads from a file-like object; use io.BytesIO in tests
    from io import BytesIO
    from docx import Document
    doc = Document()
    doc.add_paragraph("张三")
    doc.add_paragraph("软件工程专业")
    doc.add_paragraph("电子邮件: zhangsan@example.com")
    doc.add_paragraph("电话: 13800138000")
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture
def mock_resume_llm_client():
    """Return a MagicMock that simulates a successful DeepSeek resume parse response."""
    client = MagicMock()
    mock_message = MagicMock()
    mock_message.content = json.dumps({
        "name": "张三",
        "education_level": "本科",
        "contact": {"phone": "13800138000", "email": "zhangsan@example.com"},
        "education": [{"school": "清华大学", "major": "软件工程", "gpa": "3.8", "start_year": "2020", "end_year": "2024"}],
        "professional_skills": {"core": ["Python", "数据结构"], "soft": ["团队协作"], "tools": ["Git", "Docker"]},
        "certificates": {"required": [], "preferred": ["计算机二级"]},
        "experience": {
            "internship": [{"company": "字节跳动", "position": "后端实习生", "duration": "3个月", "description": "参与推荐系统开发"}],
            "projects": [{"name": "校园二手平台", "role": "技术负责人", "duration": "6个月", "description": "使用Flask构建"}],
            "extracurriculars": [{"activity": "程序设计竞赛", "role": "参赛者", "duration": "1个月", "description": "获得省级二等奖"}]
        },
        "innovation": 4.0,
        "learning": 4.5,
        "stress_resistance": 3.5,
        "communication": 4.0,
        "missing_fields": [],
        "parse_attempts": 1
    })
    mock_message.role = "assistant"
    mock_choice = MagicMock()
    mock_choice.message = mock_message
    mock_completion = MagicMock()
    mock_completion.choices = [mock_choice]
    mock_completion.model = "deepseek-chat"
    client.chat.completions.create.return_value = mock_completion
    return client


@pytest.fixture
def client(mock_deepseek_client):
    """Return a TestClient for the FastAPI app."""
    from app.main import app
    from unittest.mock import patch
    # Patch get_deepseek_client to return our mock
    with patch("app.clients.deepseek.get_deepseek_client", return_value=mock_deepseek_client):
        from fastapi.testclient import TestClient
        with TestClient(app) as c:
            yield c


@pytest.fixture
def resume_client(mock_resume_llm_client):
    """Return a TestClient with mock_resume_llm_client injected for /resume tests."""
    from app.main import app
    from app.clients.deepseek import get_deepseek_client
    from fastapi.testclient import TestClient
    app.dependency_overrides[get_deepseek_client] = lambda: mock_resume_llm_client
    try:
        with TestClient(app) as c:
            yield c
    finally:
        app.dependency_overrides.clear()
